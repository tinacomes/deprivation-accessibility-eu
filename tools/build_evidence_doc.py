#!/usr/bin/env python3
"""Build docs/evidence/depacc_results_evidence.docx from the paper pack.

Every table in the document is computed at build time from
docs/paper-pack/data/*.csv and every figure is embedded from
docs/paper-pack/figures/, so the document cannot drift from the pack
data. The short claim paragraphs are condensed from
docs/results-headlines.md (which tools/audit_paper_pack.py checks
against the same CSVs); when the headlines change, regenerate this
document and update the claim strings below in the same commit.

Usage:  python tools/build_evidence_doc.py
Deps:   pip install python-docx
"""

from __future__ import annotations

import csv
import datetime as dt
import math
import statistics
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "docs" / "paper-pack" / "data"
FIGS = ROOT / "docs" / "paper-pack" / "figures"
OUT = ROOT / "docs" / "evidence" / "depacc_results_evidence.docx"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
HEADER_FILL = "DCE6F1"
NOTE_GREY = RGBColor(0x59, 0x59, 0x59)


# ---------------------------------------------------------------- helpers
def read(name: str) -> list[dict]:
    with open(DATA / name, newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def f(x, nd=3):
    """Format a number; '' for missing."""
    if x is None or x == "":
        return ""
    v = float(x)
    return f"{v:.{nd}f}"


def fp(x):
    """Format a p-value."""
    if x is None or x == "":
        return ""
    v = float(x)
    if v < 1e-4:
        return f"{v:.1e}"
    return f"{v:.4f}".rstrip("0").rstrip(".") if v < 0.001 else f"{v:.3f}"


def fpop(x):
    return f"{int(round(float(x))):,}"


def git_state() -> str:
    try:
        h = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"], cwd=ROOT,
            capture_output=True, text=True, check=True).stdout.strip()
        b = subprocess.run(
            ["git", "rev-parse", "--abbrev-ref", "HEAD"], cwd=ROOT,
            capture_output=True, text=True, check=True).stdout.strip()
        return f"{b} @ {h}"
    except Exception:
        return "unknown"


def set_cell_bg(cell, hexfill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(shd)


class Doc:
    def __init__(self):
        self.d = Document()
        sec = self.d.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)  # A4
        for side in ("top", "bottom", "left", "right"):
            setattr(sec, f"{side}_margin", Cm(2.0))
        st = self.d.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10)
        for lvl in (1, 2, 3):
            hs = self.d.styles[f"Heading {lvl}"]
            hs.font.color.rgb = ACCENT
            hs.font.name = "Calibri"

    def h(self, text, lvl=1):
        self.d.add_heading(text, level=lvl)

    def p(self, text="", bold=False, italic=False, size=None, grey=False):
        par = self.d.add_paragraph()
        run = par.add_run(text)
        run.bold, run.italic = bold, italic
        if size:
            run.font.size = Pt(size)
        if grey:
            run.font.color.rgb = NOTE_GREY
        return par

    def note(self, text):
        par = self.p(text, italic=True, size=8.5, grey=True)
        par.paragraph_format.space_after = Pt(8)

    def bullets(self, items):
        for it in items:
            self.d.add_paragraph(it, style="List Bullet")

    def fig(self, path: Path, caption: str, width_cm=16.5):
        par = self.d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(path), width=Cm(width_cm))
        cap = self.p(caption, italic=True, size=8.5, grey=True)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)

    def table(self, header, rows, font_pt=8, caption=None):
        if caption:
            par = self.p(caption, bold=True, size=9)
            par.paragraph_format.space_after = Pt(2)
        t = self.d.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, name in enumerate(header):
            cell = t.rows[0].cells[j]
            cell.text = str(name)
            set_cell_bg(cell, HEADER_FILL)
        for row in rows:
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
        for r in t.rows:
            for c in r.cells:
                for par in c.paragraphs:
                    par.paragraph_format.space_after = Pt(0)
                    for run in par.runs:
                        run.font.size = Pt(font_pt)
                        if r is t.rows[0]:
                            run.font.bold = True
        self.p("", size=6)
        return t

    def page_break(self):
        self.d.add_page_break()


# ---------------------------------------------------------------- grades
def coverage_grades():
    """desert = main-pass outlier group; partial = peeled small cluster."""
    desert = read("cluster_null_peeled.csv")[0]["removed"].split(";")
    peeled = read("cityvector_clustered_peeled.csv")
    from collections import Counter

    counts = Counter(r["cluster_kmeans"] for r in peeled)
    small_label = min(counts, key=counts.get)
    partial = [r["city"] for r in peeled if r["cluster_kmeans"] == small_label]
    return set(desert), set(partial)


def build() -> None:
    doc = Doc()
    cityplane = read("cityplane.csv")
    desc = read("cities_descriptives.csv")
    desert, partial = coverage_grades()

    def grade(city):
        return ("desert" if city in desert
                else "partial" if city in partial else "covered")

    # ------------------------------------------------------------ title
    doc.p()
    par = doc.p("Deprivation accessibility in European cities",
                bold=True, size=22)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par = doc.p("Evidence document — headline results with tables and figures",
                size=14)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par = doc.p(
        f"Built {dt.date.today().isoformat()} from docs/paper-pack/ "
        f"({git_state()}) by tools/build_evidence_doc.py", size=9, grey=True)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.p()
    doc.h("Status and provenance", 2)
    n_countries = len({r["country"] for r in desc})
    doc.bullets([
        f"Sample complete and final: {len(desc)} functional urban areas in "
        f"{n_countries} countries (F.2 stratified draw, 4 macro-regions × 4 "
        "size strata), routed with r5 over OSM networks — walk for everyday "
        "services, car for emergency care.",
        "Final state produced by batch runs 32012438178 + 32055983121 on the "
        "depacc-results branch; every table below is recomputed from "
        "docs/paper-pack/data/ at build time.",
        "Emergency deprivation is reported in multiples of g(15 min) — 1.0 = "
        "the deprivation cost of arriving at the 15-minute EMS upper target. "
        "The re-anchoring from the superseded g(45 min) unit (factor 6.7307) "
        "was verified to machine precision on the artifact-seeded cities: "
        "identical ranks, Ginis and ratios.",
        "EMS benchmark structure used throughout: negligible cost at the "
        "3–4-minute ideal, the widely used (contested) 8-minute urban "
        "response standard, the 15-minute upper target as the real cut-off. "
        "Beyond 15 minutes no benchmark exists; conclusions that depend on "
        "the beyond-cut-off escalation form are reported conditionally (H2).",
        "Prose-vs-table consistency of the paper pack is checked by "
        "tools/audit_paper_pack.py (runs in CI).",
    ])
    doc.page_break()

    # ------------------------------------------------------------ sample
    doc.h("Sample — 67 cities at a glance")
    doc.p(
        "One row per FUA (sorted by population). Emergency means are in "
        "multiples of g(15 min); HH50 is the compounding population share at "
        "the median split; grade is the emergency-coverage severity ordering "
        "of H4 (covered → partial desert → desert; the covered/partial "
        "boundary is fuzzy).")
    rows = []
    for r in sorted(desc, key=lambda x: -float(x["population"])):
        rows.append([
            r["name"], r["country"], r["region"], fpop(r["population"]),
            r["n_emergency_services"], grade(r["city"]),
            f(r["mean_everyday"], 3), f(r["mean_emergency"], 2),
            f(r["gini_everyday"], 3), f(r["gini_emergency"], 3),
            f(r["spearman_rho"], 2), f(r["compounding_pop_share_50"], 2),
        ])
    doc.table(
        ["City", "CC", "Region", "Population", "EMS srv", "Grade",
         "mean ev", "mean em", "Gini ev", "Gini em", "ρ", "HH50"],
        rows, font_pt=7,
        caption="Table S1 — city descriptives (cities_descriptives.csv)")
    doc.note(
        "EMS srv = number of emergency service types present (8 cities have "
        "a single one — their emergency composite rests on one facility "
        "layer; limitation noted in methods §3).")
    doc.page_break()

    # ------------------------------------------------------------ H1
    doc.h("H1 — City size buys everyday access, but not emergency care")
    doc.p(
        "Mean everyday deprivation falls ~19 % per e-fold of population; "
        "mean emergency deprivation shows no detectable size gradient, and "
        "the TOST bounds it within ±0.13 (equivalence at ±0.10 not shown — "
        "say “no detectable scaling, bounded within ±0.13”, never “proven "
        "flat”). The difference between the two slopes is itself "
        "significant.")
    clus = {r["outcome"]: r for r in read("inference_scaling_clustered.csv")}
    rows = []
    for out in ("mean_everyday", "mean_emergency",
                "gini_everyday", "gini_emergency"):
        r = clus[out]
        rows.append([out, f(r["elasticity"]), f(r["se_cluster_country"]),
                     fp(r["p_cluster_country"]),
                     fp(r["p_wild_cluster_bootstrap"]),
                     r["n_cities"], r["n_countries"]])
    doc.table(
        ["Outcome", "Elasticity", "SE (country)", "p (clustered)",
         "p (wild bootstrap)", "n cities", "n countries"], rows,
        caption="Table H1a — size elasticities, country-clustered "
                "(inference_scaling_clustered.csv)")
    paired = {r["measure"]: r for r in read("inference_regime_paired.csv")}
    rows = [[m, f(paired[m]["gradient_everyday"]),
             f(paired[m]["gradient_difference_emergency"]),
             f(paired[m]["se_cluster_country"]),
             fp(paired[m]["p_wild_cluster_bootstrap"])]
            for m in ("mean", "gini")]
    doc.table(
        ["Measure", "Everyday gradient", "Emergency − everyday",
         "SE (country)", "p (wild bootstrap)"], rows,
        caption="Table H1b — paired regime slope difference "
                "(inference_regime_paired.csv)")
    eq = read("inference_equivalence.csv")[0]
    doc.table(
        ["Outcome", "Elasticity", "TOST bound", "p (TOST)",
         "Smallest passing bound", "Conclusion"],
        [[eq["outcome"], f(eq["elasticity"]), eq["equivalence_bound"],
          fp(eq["p_tost"]), f(eq["smallest_passing_bound"], 3),
          eq["conclusion"]]],
        caption="Table H1c — equivalence test (inference_equivalence.csv)")
    infl = read("inference_influence.csv")
    rows = [[r["outcome"], f(r["elasticity"]),
             f"{f(r['loo_min'])} … {f(r['loo_max'])}", r["loo_worst_city"],
             f(r["elasticity_huber"]), f(r["elasticity_median_reg"]),
             f(r["elasticity_excl_desert"]), fp(r["p_excl_desert"])]
            for r in infl]
    doc.table(
        ["Outcome", "Elasticity", "LOO envelope", "Worst-case city",
         "Huber", "Median reg.", "Excl. deserts", "p excl."], rows,
        caption="Table H1d — influence and robustness "
                "(inference_influence.csv)")
    doc.fig(FIGS / "scaling_elasticity.png",
            "Figure H1 — scaling_elasticity.png: log-log size gradients by "
            "regime.")
    doc.page_break()

    # ------------------------------------------------------------ H2
    doc.h("H2 — Everyday inequality grows with size; "
          "emergency inequality does not")
    doc.p(
        "The everyday-deprivation Gini rises with population under all 14 "
        "deprivation parameterisations; the emergency Gini is size-flat at "
        "baseline. The emergency half is conditional on beyond-benchmark "
        "escalation and must always be reported as such: the EMS benchmarks "
        "define the cost scale up to the 15-minute cut-off and all three "
        "literature-grounded escalation forms coincide at the 4/8/15-minute "
        "anchors; beyond the cut-off no benchmark exists, and that is "
        "exactly where the forms diverge.")
    spec = read("specification_curve.csv")
    ev = sorted(float(r["elasticity"]) for r in spec
                if r["outcome"] == "gini_everyday" and r["elasticity"])
    ev_sig = sum(1 for r in spec
                 if r["outcome"] == "gini_everyday" and r["p"]
                 and float(r["p"]) < 0.05)
    doc.bullets([
        f"Everyday Gini slope across all {len(ev)} parameterisations: "
        f"{ev[0]:+.3f} to {ev[-1]:+.3f}, significant in {ev_sig}/{len(ev)} "
        "(the exception is the concave Box-Cox everyday form swap — a "
        "caveat to state).",
        "Saturating escalation — the bounded survival-based curve "
        "(1 = full deprivation; saturation at ~30–45 min emerging from "
        "the benchmark anchors): emergency Gini size-flat — the "
        "“1 = full deprivation” framing gives the same headline as the "
        "baseline.",
        "Polynomial (Box-Cox) beyond-cut-off escalation — the "
        "Cantillo/Delgado-Lindeman ambulance-DCF line: emergency Gini "
        "size-flat, in agreement with the deprivation-free travel-time "
        "Gini (a benchmark-free anchor).",
        "Exponential escalation — the Holguín-Veras line: emergency Gini "
        "rising with size. Verified mechanism: beyond-cut-off tail shares "
        "are size-flat, but exponential pricing makes the Gini track "
        "extreme-tail depth — it saturates near its ceiling (median 0.96) "
        "and decouples from the travel-time Gini (ρ = 0.04).",
        "Safe unconditional statements: everyday inequality rises under "
        "everything; emergency inequality never falls with size; flat vs "
        "rising is decided by the escalation form, which no benchmark "
        "pins down.",
    ])
    rows = []
    for r in spec:
        if r["outcome"] != "gini_emergency" or not r["elasticity"]:
            continue
        if r["variant"] in ("baseline", "emergency_lam1.4",
                            "emergency_lam2.2",
                            "formswap_emergency_survival",
                            "formswap_emergency_exponential"):
            label = {
                "baseline": "Box-Cox λ 1.8 (baseline)",
                "emergency_lam1.4": "Box-Cox λ 1.4",
                "emergency_lam2.2": "Box-Cox λ 2.2",
                "formswap_emergency_survival":
                    "Saturating (survival-based, bounded, 1 = full)",
                "formswap_emergency_exponential":
                    "Exponential (Holguín-Veras line)",
            }[r["variant"]]
            rows.append([label, f(r["elasticity"]), f(r["se"]), fp(r["p"])])
    doc.table(
        ["Beyond-cut-off escalation form", "Emergency-Gini elasticity",
         "SE", "p"], rows,
        caption="Table H2a — the conditional emergency-Gini result "
                "(specification_curve.csv)")
    gc = [r for r in read("deprivation_vs_access.csv")
          if r["block"] == "gini_correlation"]
    doc.table(
        ["Regime", "Pearson r (dep-Gini vs time-Gini)", "Spearman ρ"],
        [[r["regime"], f(r["elasticity"]), f(r["p"])] for r in gc],
        caption="Table H2b — deprivation Gini vs deprivation-free "
                "travel-time Gini (deprivation_vs_access.csv)")
    doc.fig(FIGS / "deprivation_curves.png",
            "Figure H2 — deprivation_curves.png: everyday level curve; the "
            "emergency benchmark window (0–15 min), where all three "
            "escalation forms coincide at the EMS anchors by construction; "
            "and the beyond-cut-off region (log scale), where they diverge "
            "— saturating 1.6× / Box-Cox 11× / exponential 120× the "
            "benchmark cost at 60 min. The dotted line is the linear loss "
            "a pure-access minutes average implies.")
    doc.page_break()

    # ------------------------------------------------------------ H3
    doc.h("H3 — The divergence has a geography, and it is country-level")
    doc.p(
        "Regional contrasts tested with the honest unit: whole countries "
        "permuted across macro-regions (cities nest in countries). CEE "
        "compounds — highest emergency inequality and strongest "
        "everyday–emergency coupling; the West's inequality is "
        "everyday-side, the North's emergency-side. The everyday Gini's "
        "apparent regional contrast does not survive country-level testing "
        "— do not claim it.")
    rows = []
    for r in read("inference_regional.csv"):
        rows.append([r["outcome"], f(r["cohen_f_city"], 2),
                     fp(r["p_city_anova_anticonservative"]),
                     fp(r["p_permutation_countries"]),
                     fp(r["p_mixed_model"])])
    doc.table(
        ["Outcome", "Cohen f", "p city-ANOVA (anti-conservative)",
         "p permutation (countries)", "p mixed model"], rows,
        caption="Table H3 — regional contrasts (inference_regional.csv); "
                "the country permutation is the primary test")
    doc.fig(FIGS / "region_strips.png",
            "Figure H3 — region_strips.png: per-region strips of the key "
            "indicators.")
    doc.page_break()

    # ------------------------------------------------------------ H4
    doc.h("H4 — No city types; regional gradients plus one desert group")
    main = read("cluster_null.csv")[0]
    peel = read("cluster_null_peeled.csv")[0]
    fr = read("cluster_feature_robustness.csv")[0]
    doc.p(
        "k-means cuts 62-vs-5: the emergency-desert capitals "
        f"({', '.join(sorted(desert))}). Real split, but an outlier group, "
        "not a typology (no agreement with the region partition); Ward "
        "nests inside it — a severity ordering, not two algorithms "
        "agreeing on one group. The peeled re-clustering (outliers removed, "
        "everything recomputed) finds one further split along the same axis "
        f"at lower intensity: {len(partial)} partial-desert cities. The "
        "citable reading: the only recoverable discrete structure is a "
        "severity ordering on one axis — emergency-periphery coverage "
        "(covered → partial desert → desert) — a gradient made visible, "
        "not city types; the covered/partial boundary is explicitly fuzzy.")
    doc.table(
        ["Pass", "n", "k", "Silhouette", "Bootstrap ARI", "Gaussian-null p",
         "Small group"],
        [["Main", "67", main["k"], f(main["observed_silhouette"], 2),
          f(main["stability_ari"], 2), fp(main["p_null_gaussian"]),
          f"5 deserts: {', '.join(sorted(desert))}"],
         ["Peeled", peel["n_cities"], peel["k"], f(peel["silhouette"], 2),
          f(peel["stability_ari"], 2), fp(peel["p_null_gaussian"]),
          f"{len(partial)} partial deserts: {', '.join(sorted(partial))}"]],
        caption="Table H4a — clustering diagnostics (cluster_null.csv, "
                "cluster_null_peeled.csv)")
    doc.table(
        ["Feature set", "n features", "ARI vs published (main)",
         "Same 5 capitals", "ARI vs published (peeled)"],
        [[fr["feature_set"], fr["n_features"], f(fr["ari_vs_published"], 2),
          fr["small_group_matches_published"],
          f(fr["peeled_ari_vs_published"], 2)]],
        caption="Table H4b — feature-set robustness: full pipeline re-run "
                "with the 8/15-min benchmark shares included "
                "(cluster_feature_robustness.csv)")
    doc.note(
        "The extended feature set flags Brăila+Łomża as their own "
        "micro-group — converging evidence that they sit on the "
        "covered/partial boundary rather than in any type; the same two "
        "cities moved across that boundary between extraction states "
        "(50/12 → 48/14). Ordering robust, boundary fuzzy.")
    doc.fig(FIGS / "cityplane.png",
            "Figure H4 — cityplane.png: the 67 cities in the "
            "everyday-vs-emergency inequality plane, coloured by "
            "macro-region; the emergency-desert group is circled and the "
            "grey whiskers are the deprivation-curvature envelopes.")
    doc.page_break()

    # ------------------------------------------------------------ H5
    doc.h("H5 — Who carries it: children, almost everywhere")
    doc.p(
        "Census-harmonised strata pooled across all 67 cities against the "
        "covered-cell reference. Children live in more-deprived and "
        "higher-compounding cells in ~88 % of cities; the elderly do not "
        "compound at the harmonised European level (below parity across "
        "most of CEE) — sharper Tier-2 national results (e.g. Hamburg) use "
        "different definitions and are reported per city, never pooled.")
    rows = []
    for r in read("vulnerability_summary.csv"):
        rows.append([
            r["stratum"], r["n_cities"],
            f(r["median_everyday_ratio_covered"], 2),
            f"{100 * float(r['share_cities_everyday_ratio_gt1']):.0f} %",
            f(r["median_emergency_ratio_covered"], 2),
            f"{100 * float(r['share_cities_emergency_ratio_gt1']):.0f} %",
            f(r["median_hh_share_gap_covered"], 3),
            f"{100 * float(r['share_cities_hh_gap_gt0']):.0f} %"])
    doc.table(
        ["Stratum", "n", "Median everyday ratio", "> parity",
         "Median emergency ratio", "> parity", "Median HH gap",
         "> parity"], rows,
        caption="Table H5 — pooled stratum ratios, covered-cell reference "
                "(vulnerability_summary.csv)")
    doc.fig(FIGS / "vulnerability_strata.png",
            "Figure H5 — vulnerability_strata.png: per-city stratum ratios.")
    doc.page_break()

    # ------------------------------------------------------------ H6
    doc.h("H6 — Compounding is the norm, not the exception")
    hh = [float(r["compounding_pop_share_50"]) for r in cityplane]
    rho = [float(r["spearman_rho"]) for r in cityplane]
    n_above = sum(1 for v in hh if v > 0.25)
    sens = read("deprivation_sensitivity_summary.csv")
    med_w = {}
    for axis in ("curvature", "form_swap", "threshold"):
        w = [float(r["width"]) for r in sens
             if r["axis"] == axis and r["target"] == "share_HH"]
        med_w[axis] = statistics.median(w) if w else math.nan
    doc.p(
        f"The HH share at the median split averages {100 * statistics.mean(hh):.0f} % "
        f"against the 25 % independence benchmark and exceeds it in "
        f"{n_above} of {len(hh)} cities; coupling ρ spans "
        f"{min(rho):.2f} to {max(rho):.2f}. The class shares barely notice "
        "the deprivation function: median per-city envelope width "
        f"{100 * med_w['curvature']:.1f} pp under the curvature grid and "
        f"{100 * med_w['form_swap']:.1f} pp under the form swap, against "
        f"{100 * med_w['threshold']:.1f} pp for the “how high is high” "
        "threshold — which is exactly why the continuous "
        "compounding_intensity is the headline and the class shares carry "
        "a threshold sweep.")
    doc.table(
        ["Sensitivity axis", "Median per-city envelope width of the HH "
         "share", "Cities"],
        [["Deprivation curvature grid",
          f"{100 * med_w['curvature']:.1f} pp", str(len(hh))],
         ["Functional-form swap", f"{100 * med_w['form_swap']:.1f} pp",
          str(len(hh))],
         ["“How high is high” threshold",
          f"{100 * med_w['threshold']:.1f} pp", str(len(hh))]],
        caption="Table H6 — what moves the compounding classes "
                "(deprivation_sensitivity_summary.csv)")
    doc.fig(FIGS / "compounding_gallery.png",
            "Figure H6a — compounding_gallery.png: co-location typologies "
            "across the sample.")
    doc.fig(FIGS / "rho_ranked.png",
            "Figure H6b — rho_ranked.png: everyday–emergency coupling ρ, "
            "ranked.")

    # ------------------------------------------------------------ H7
    doc.h("H7 — Deprivation is not access: the claims survive, "
          "the framing earns its keep")
    doc.p(
        "Re-running the size regressions on plain minutes: the claims "
        "survive without any deprivation function; the deprivation outcome "
        "is better behaved (higher R² — the loss function discounts "
        "variation in the flat part of the curve that carries no welfare "
        "content); and the deserts are invisible to access averages — only "
        "the convex, clinically anchored DCF prices the tail that puts "
        "them there.")
    rows = []
    for r in read("deprivation_vs_access.csv"):
        if r["block"] != "size_elasticity":
            continue
        rows.append([r["regime"], r["kind"], r["outcome"],
                     f(r["elasticity"]), f(r["se"]), fp(r["p"]),
                     f(r["r2"], 2)])
    doc.table(
        ["Regime", "Kind", "Outcome", "Elasticity", "SE", "p", "R²"], rows,
        caption="Table H7a — deprivation vs pure-access size elasticities "
                "(deprivation_vs_access.csv)")
    rows = []
    for r in read("desert_access_contrast.csv"):
        rows.append([r["city"], r["country"],
                     f(r["emergency_median_time_min"], 1),
                     f"{float(r['median_time_ratio_vs_sample']):.2f}×",
                     f(r["emergency_p90_time_min"], 1),
                     f(r["mean_emergency_deprivation"], 2),
                     f"{float(r['deprivation_ratio_vs_sample']):.2f}×"])
    doc.table(
        ["City", "CC", "Median EMS min", "vs sample", "p90 EMS min",
         "Mean emergency deprivation", "vs sample"], rows,
        caption="Table H7b — the deserts are invisible to access medians "
                "(desert_access_contrast.csv)")
    rows = []
    for r in read("scaling_by_grade.csv"):
        rows.append([r["grade"], r["n_cities"], f(r["elasticity"]),
                     f(r["se"]), fp(r["p"])])
    doc.table(
        ["Coverage grade", "n", "Within-grade elasticity", "SE", "p"], rows,
        caption="Table H7c — size gradient within coverage grades "
                "(scaling_by_grade.csv)")
    doc.note(
        "The grade LEVEL shift is the result (median mean-emergency "
        "deprivation 1.02 / 1.87 / 3.37 across covered / partial / desert; "
        "country-clustered Wald p ≈ 1e-15). The slope × grade interaction "
        "is boundary-sensitive (p = 0.010 under 48/14/5, p = 0.17 under "
        "the pre-refresh 50/12/5) — lead with “coverage class, not city "
        "size, sets emergency deprivation” and never headline a "
        "grade-specific size gradient.")
    doc.fig(FIGS / "scaling_by_coverage_grade.png",
            "Figure H7 — scaling_by_coverage_grade.png.")
    doc.page_break()

    # ------------------------------------------------------- robustness
    doc.h("Robustness inventory")
    doc.p(
        "Everything the claims were pushed through, in one place: the "
        "specification curve (14 deprivation parameterisations), rank "
        "agreement of the city orderings, the per-city assumption "
        "footprint (flip cells), and the level envelopes.")
    ra = [r for r in read("rank_agreement.csv") if r["variant"]]
    curv = [float(r["spearman_rho"]) for r in ra
            if not r["variant"].startswith("formswap_")]
    ev_swap = [float(r["spearman_rho"]) for r in ra
               if r["variant"] == "formswap_everyday_box_cox"]
    def _em_swap_rho(variant):
        rows_ = [r for r in ra if r["variant"] == variant
                 and r["target"] == "gini_emergency"]
        return float(rows_[0]["spearman_rho"]) if rows_ else None

    exp_rho = _em_swap_rho("formswap_emergency_exponential")
    sv_rho = _em_swap_rho("formswap_emergency_survival")
    bullets = [
        f"Rank agreement: minimum Spearman ρ {min(curv):.2f} over the "
        f"curvature grid and {min(ev_swap):.2f} under the everyday form "
        "swap — the rank-based results are parameterisation-free in that "
        "scope.",
    ]
    if exp_rho is not None:
        bullets.append(
            "Scoped exception (a change of measurand, not noise): under "
            "exponential escalation the emergency-Gini ordering becomes "
            f"an extreme-tail-depth ordering (ρ = {exp_rho:.2f} vs "
            "baseline).")
    if sv_rho is not None:
        bullets.append(
            "Scoped exception: under the bounded survival curve it "
            "becomes a benchmark-window inequality ordering "
            f"(ρ = {sv_rho:.2f} vs baseline) — the tail beyond saturation "
            "no longer distinguishes cities.")
    doc.bullets(bullets)
    fc = [float(r["sensitive_pop_share"]) for r in read("flip_cells.csv")]
    doc.table(
        ["Metric", "Value"],
        [["Population-weighted sensitive share, mean over 67 cities",
          f"{100 * statistics.mean(fc):.1f} %"],
         ["Range", f"{100 * min(fc):.1f} – {100 * max(fc):.1f} %"],
         ["Complementary stable share (keeps its class under every "
          "deprivation parameterisation)",
          f"{100 * (1 - statistics.mean(fc)):.1f} %"]],
        caption="Table R1 — flip cells: the spatial footprint of the "
                "deprivation assumption (flip_cells.csv); distinct from "
                "the ~24 % routing-engine flip (E.1)")
    rows = []
    for axis in ("curvature", "form_swap"):
        for target in ("gini_everyday", "gini_emergency", "divergence_gap"):
            w = [float(r["width"]) for r in sens
                 if r["axis"] == axis and r["target"] == target]
            if w:
                rows.append([axis, target, f(statistics.median(w), 3),
                             f"{f(min(w), 3)} – {f(max(w), 3)}",
                             str(len(w))])
    doc.table(
        ["Axis", "Target", "Median envelope width", "Range", "Cities"],
        rows,
        caption="Table R2 — level envelopes: the Gini LEVELS move a great "
                "deal under the same sweep (the calibration carries "
                "information); the slope claims are defended by the whole "
                "specification curve, not by any one parameterisation "
                "(deprivation_sensitivity_summary.csv)")
    doc.fig(FIGS / "specification_curve.png",
            "Figure R1 — specification_curve.png: the headline slopes "
            "under every deprivation parameterisation.")
    doc.page_break()

    # ------------------------------------------------- service table
    doc.h("Service-level accessibility (pooled)")
    doc.p(
        "Deprivation-free per-service accessibility, pooled over cities "
        "(median [min–max] across cities). The 15-minute rows connect to "
        "the 15-minute-city literature; emergency rows carry the "
        "free-flow-driving caveat (methods §3).")
    rows = []
    for r in read("accessibility_by_service_pooled.csv"):
        rows.append([
            r["regime"], r["service"], r["n_cities"],
            f"{f(r['n_facilities_median'], 0)} "
            f"[{f(r['n_facilities_min'], 0)}–{f(r['n_facilities_max'], 0)}]",
            f"{f(r['pop_median_time_min_median'], 0)} "
            f"[{f(r['pop_median_time_min_min'], 0)}–"
            f"{f(r['pop_median_time_min_max'], 0)}]",
            f"{100 * float(r['pop_share_beyond_15min_median']):.0f} % "
            f"[{100 * float(r['pop_share_beyond_15min_min']):.0f}–"
            f"{100 * float(r['pop_share_beyond_15min_max']):.0f}]",
            f"{100 * float(r['pop_share_beyond_30min_median']):.1f} % "
            f"[{100 * float(r['pop_share_beyond_30min_min']):.1f}–"
            f"{100 * float(r['pop_share_beyond_30min_max']):.1f}]"])
    doc.table(
        ["Regime", "Service", "n cities", "Facilities med [range]",
         "Median time med [range]", "> 15 min med [range]",
         "> 30 min med [range]"], rows, font_pt=7,
        caption="Table S2 — per-service accessibility "
                "(accessibility_by_service_pooled.csv)")
    doc.page_break()

    # ------------------------------------------------- exemplar maps
    doc.h("Exemplar city maps")
    doc.p(
        "Four of the 67 per-city compounding maps (median-split "
        "co-location typology), one per coverage situation. The full "
        "gallery with a grade/HH/ρ index is in "
        "docs/paper-pack/figures/cities/ (README.md). Read these for "
        "spatial pattern, not per-cell class: ~24 % of population changes "
        "class between routing engines (E.1).")
    exemplars = [
        ("hamburg", "Hamburg (DE) — covered, Tier-2 deep-dive city"),
        ("paris", "Paris (FR) — covered, largest FUA in the sample"),
        ("bucuresti", "București (RO) — emergency-desert capital"),
        ("lomza", "Łomża (PL) — boundary city (covered ↔ partial desert)"),
    ]
    for city, caption in exemplars:
        p = FIGS / "cities" / f"{city}.png"
        if p.exists():
            doc.fig(p, f"{caption} — figures/cities/{city}.png",
                    width_cm=14.5)

    # ------------------------------------------------- not claimed
    doc.h("What is deliberately NOT claimed")
    doc.bullets([
        "Absolute travel times and Gini levels under the friction engine "
        "(E.1: levels understated 34–314 %; r5 is the primary engine).",
        "Per-cell typology classes on any map (~24 % engine flip; maps are "
        "for spatial patterns).",
        "A causal or longitudinal reading — every gradient is "
        "cross-sectional space-for-time.",
        "An everyday-Gini regional contrast (H3) and elderly compounding "
        "at the European level (H5).",
        "An unconditional emergency-Gini trend (H2) — always report the "
        "escalation-form conditionality; safe statements: “never falls "
        "with size”, “the everyday Gini rises under everything”.",
        "Rank-agreement universality — under exponential escalation the "
        "emergency-Gini ranking is a different measurand (reported as "
        "such, with the saturation numbers).",
        "Two clustering algorithms agreeing on one desert group (Ward "
        "nests inside k-means; corroboration of a severity ordering).",
    ])
    doc.h("Known limitations", 2)
    doc.bullets([
        "GP-density artifact in the everyday composite for a few cities "
        "(OSM completeness; flagged in methods).",
        "8 cities carry a single emergency service type — their emergency "
        "composite rests on one facility layer (Table S1).",
        "Emergency routing uses free-flow car speeds — response times are "
        "optimistic; the EMS anchors are applied to the same scale on both "
        "sides so comparisons are internally consistent (methods §3).",
        "DE registry facility counts flagged for re-verification before "
        "publication (E.2 note).",
        "Every cross-city gradient is cross-sectional space-for-time; no "
        "longitudinal claim is made.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.d.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
